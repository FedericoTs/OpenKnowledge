"""Word documents.

python-docx exposes paragraphs and tables as separate object models, which is
better than it sounds: the alternative approach - convert to HTML and flatten -
exists precisely because most extractors lose tables, and here they never need
recovering.

The one thing the object model does not give you is document order. ``.paragraphs``
and ``.tables`` are two flat lists, so a table appears at the end rather than
where it sits in the text. Walking the underlying body XML restores the real
sequence, which matters because a table's meaning usually comes from the heading
above it.
"""

from __future__ import annotations

import logging

from .blocks import (
    Block,
    BlockKind,
    ParsedDocument,
    looks_like_header_row,
    normalise,
    table_row_text,
)

log = logging.getLogger(__name__)


def _heading_level(style_name: str) -> int:
    """Heading depth from a Word style name, or 0 if it is not a heading."""
    name = (style_name or "").strip().lower()
    if name in ("title", "subtitle"):
        return 1
    if name.startswith("heading"):
        digits = "".join(ch for ch in name if ch.isdigit())
        return int(digits) if digits else 1
    return 0


def _is_list(style_name: str) -> bool:
    name = (style_name or "").strip().lower()
    return "list" in name or "bullet" in name


def parse_docx(data: bytes, *, title: str | None = None) -> ParsedDocument:
    """Parse a .docx into blocks, in document order, tables intact."""
    try:
        import docx
        from docx.document import Document as DocxDocument
        from docx.table import Table
        from docx.text.paragraph import Paragraph
    except ImportError:  # pragma: no cover - dependency is declared
        return ParsedDocument(warnings=("python-docx is not installed",))

    import io

    try:
        document = docx.Document(io.BytesIO(data))
    except Exception as exc:  # noqa: BLE001 - malformed or encrypted file
        log.warning("could not read .docx: %s", exc)
        return ParsedDocument(warnings=(f"could not read this Word document: {exc}",))

    blocks: list[Block] = []
    heading_path: list[str] = []
    doc_title = title

    def push_heading(text: str, level: int) -> None:
        nonlocal doc_title
        del heading_path[level - 1 :]
        heading_path.append(text)
        if doc_title is None and level == 1:
            doc_title = text
        blocks.append(
            Block(
                kind=BlockKind.HEADING,
                text=text,
                heading_path=tuple(heading_path[:-1]),
                level=level,
            )
        )

    for item in _in_document_order(document, DocxDocument, Paragraph, Table):
        if isinstance(item, Paragraph):
            text = normalise(item.text)
            if not text:
                continue
            style = item.style.name if item.style is not None else ""
            level = _heading_level(style)
            if level:
                push_heading(text, level)
            else:
                blocks.append(
                    Block(
                        kind=BlockKind.LIST_ITEM if _is_list(style) else BlockKind.PARAGRAPH,
                        text=text,
                        heading_path=tuple(heading_path),
                    )
                )
        else:
            blocks.extend(_table_blocks(item, tuple(heading_path)))

    return ParsedDocument(blocks=tuple(blocks), title=doc_title)


def _in_document_order(document: object, DocxDocument: type, Paragraph: type, Table: type):
    """Yield paragraphs and tables in the order they appear in the body.

    ``.paragraphs`` and ``.tables`` are separate lists, so relying on them puts
    every table after all the prose - detaching each one from the heading that
    gives it meaning.
    """
    body = document.element.body  # type: ignore[attr-defined]
    parent = document
    for child in body.iterchildren():
        tag = child.tag.rsplit("}", 1)[-1]
        if tag == "p":
            yield Paragraph(child, parent)
        elif tag == "tbl":
            yield Table(child, parent)


def _table_blocks(table: object, heading_path: tuple[str, ...]) -> list[Block]:
    blocks: list[Block] = []
    headers: list[str] = []
    for row in table.rows:  # type: ignore[attr-defined]
        cells = [normalise(cell.text) for cell in row.cells]
        if not headers and looks_like_header_row(cells):
            headers = cells
            continue
        text = table_row_text(headers, cells)
        if text:
            blocks.append(Block(kind=BlockKind.TABLE_ROW, text=text, heading_path=heading_path))
    return blocks
